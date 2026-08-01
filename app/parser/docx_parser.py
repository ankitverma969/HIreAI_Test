from pathlib import Path

from docx import Document
from loguru import logger

from app.core.exceptions import ParsingException
from app.parser.base_parser import BaseParser


class DocxParser(BaseParser):
    """Document parser implementation for Microsoft Word DOCX files."""

    def parse(self, file_path: Path) -> str:
        """Extracts text paragraphs and tables from DOCX files.

        Args:
            file_path: Filesystem path to DOCX.

        Returns:
            Extracted text content.

        Raises:
            ParsingException: If word file is corrupt or unreadable.
        """
        try:
            logger.info(f"Parsing DOCX document: '{file_path.name}'")
            doc = Document(str(file_path))
            text_parts = []

            # 1. Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text:
                    text_parts.append(paragraph.text)

            # 2. Extract tables content (optional layout preservation)
            for table in doc.tables:
                for row in table.rows:
                    row_cells_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_cells_text:
                        text_parts.append(" | ".join(row_cells_text))

            full_text = "\n".join(text_parts)

            if not full_text.strip():
                raise ParsingException(f"DOCX document contains no readable text: '{file_path.name}'")

            return full_text
        except Exception as e:
            if isinstance(e, ParsingException):
                raise e
            logger.error(f"Error parsing DOCX file '{file_path.name}': {str(e)}")
            raise ParsingException(f"Failed to parse DOCX document '{file_path.name}': {str(e)}") from e
