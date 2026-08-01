from pathlib import Path

import fitz  # PyMuPDF
import pytest
from docx import Document

from app.core.exceptions import ParsingException
from app.parser.parser_service import ParserService


@pytest.fixture
def mock_txt_file(tmp_path: Path) -> Path:
    """Fixture generating a plain text file."""
    file_path = tmp_path / "resume.txt"
    file_path.write_text("John Doe\nEmail: john@example.com\nSkills: Python, Go", encoding="utf-8")
    return file_path


@pytest.fixture
def mock_pdf_file(tmp_path: Path) -> Path:
    """Fixture generating a valid PDF file using PyMuPDF."""
    file_path = tmp_path / "resume.pdf"
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text(
        (50, 50),
        "Jane Smith\nEmail: jane@example.com\nPhone: +1-555-123-4567\nEducation: B.Tech 2020\nExperience: Engineer at GCP (2021-Present)"
    )
    doc.save(str(file_path))
    doc.close()
    return file_path


@pytest.fixture
def mock_docx_file(tmp_path: Path) -> Path:
    """Fixture generating a valid DOCX file using python-docx."""
    file_path = tmp_path / "resume.docx"
    doc = Document()
    doc.add_paragraph("Bob Jones")
    doc.add_paragraph("Email: bob@example.com")
    doc.add_paragraph("Experience: Developer at AWS (Jan 2020 - Dec 2022)")
    doc.save(str(file_path))
    return file_path


@pytest.mark.anyio
async def test_parse_text_file(mock_txt_file: Path) -> None:
    """Verifies that plain text files are successfully read."""
    res = await ParserService.parse_document(mock_txt_file)
    assert res["file_name"] == "resume.txt"
    assert res["pages"] == 0
    assert "john@example.com" in res["cleaned_text"]
    assert res["parser_used"] == "TextParser"


@pytest.mark.anyio
async def test_parse_pdf_file(mock_pdf_file: Path) -> None:
    """Verifies that PDF pages text and page count are extracted."""
    res = await ParserService.parse_document(mock_pdf_file)
    assert res["file_name"] == "resume.pdf"
    assert res["pages"] == 1
    assert "jane@example.com" in res["cleaned_text"]
    assert res["parser_used"] == "PDFParser"


@pytest.mark.anyio
async def test_parse_docx_file(mock_docx_file: Path) -> None:
    """Verifies that Word DOCX text paragraphs are extracted."""
    res = await ParserService.parse_document(mock_docx_file)
    assert res["file_name"] == "resume.docx"
    assert res["pages"] == 0
    assert "bob@example.com" in res["cleaned_text"]
    assert res["parser_used"] == "DocxParser"


@pytest.mark.anyio
async def test_parse_invalid_format(tmp_path: Path) -> None:
    """Verifies that unsupported formats raise ParsingException."""
    bad_file = tmp_path / "resume.jpg"
    bad_file.write_bytes(b"dummy jpeg content")

    with pytest.raises(ParsingException):
        await ParserService.parse_document(bad_file)
